"""
SCI 논문 생성기 — DOCX 내보내기
1) 원고 본문(Times New Roman 12pt, 2배 줄간격, 1인치 여백의 SCI 투고 서식 / 국문일 경우 맑은 고딕)
2) 저자용 체크리스트 리포트 (제출용 문서가 아니라 작업용 참고 문서)
3) 예상 심사위원 Q&A 리포트
"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent.parent))
from backend.config import KOREAN_DOCX_FONT

EN_BODY_FONT = "Times New Roman"
EN_REPORT_FONT = "Arial"
BODY_SIZE = Pt(12)

SECTION_LABELS = {
    "abstract": "Abstract",
    "introduction": "Introduction",
    "methods": "Methods",
    "results": "Results",
    "discussion": "Discussion",
    "results_and_discussion": "Results and Discussion",
    "conclusion": "Conclusion",
    "thematic_sections": "Thematic Review",
    "references": "References",
}

NUMBERED_KEYS_EXCLUDED = {"abstract", "references"}


def _label_for(key: str) -> str:
    return SECTION_LABELS.get(key, key.replace("_", " ").title())


def _apply_font(font_obj, font_name: str) -> None:
    """python-docx Font 객체에 라틴+동아시아(eastAsia) 폰트를 함께 지정한다.
    font.name만 설정하면 한글 등 CJK 문자에는 적용되지 않고 Word 기본 폰트로 렌더링되므로,
    w:rFonts의 eastAsia 속성도 함께 채워 한글 문서에서도 지정한 폰트가 실제로 적용되게 한다."""
    font_obj.name = font_name
    rpr = font_obj._element
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        rpr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def _body_font(language: str) -> str:
    return KOREAN_DOCX_FONT if language == "ko" else EN_BODY_FONT


def _report_font(language: str) -> str:
    return KOREAN_DOCX_FONT if language == "ko" else EN_REPORT_FONT


def _set_base_style(document: Document, body_font: str) -> None:
    style = document.styles["Normal"]
    _apply_font(style.font, body_font)
    style.font.size = BODY_SIZE
    style.paragraph_format.line_spacing = 2.0
    for section in document.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)


def _add_heading(document: Document, text: str, body_font: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(text)
    _apply_font(run.font, body_font)
    run.font.size = Pt(13)
    run.bold = True


def _add_paragraphs(document: Document, paragraphs, missing_text: str) -> None:
    if not paragraphs:
        p = document.add_paragraph()
        p.paragraph_format.line_spacing = 2.0
        run = p.add_run(missing_text)
        run.italic = True
        return
    for text in paragraphs:
        p = document.add_paragraph()
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.first_line_indent = Inches(0.3)
        p.add_run(text)


def build_manuscript_docx(draft: dict, output_path: str) -> None:
    language = draft.get("language", "en")
    body_font = _body_font(language)
    is_ko = language == "ko"

    document = Document()
    _set_base_style(document, body_font)

    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.line_spacing = 1.5
    title_run = title_p.add_run(draft.get("title") or ("제목 없음" if is_ko else "Untitled Manuscript"))
    _apply_font(title_run.font, body_font)
    title_run.bold = True
    title_run.font.size = Pt(15)

    keywords = draft.get("keywords") or []
    if keywords:
        kw_p = document.add_paragraph()
        kw_p.paragraph_format.line_spacing = 1.5
        kw_p.paragraph_format.space_before = Pt(6)
        label = kw_p.add_run("주요어: " if is_ko else "Keywords: ")
        label.bold = True
        kw_p.add_run(", ".join(keywords))

    sections = draft.get("sections", {})
    section_order = draft.get("section_order") or list(sections.keys())
    missing_text = "[NEEDS DATA: 이 섹션에 대해 생성된 내용이 없습니다]" if is_ko else "[NEEDS DATA: no content generated for this section]"

    number = 1
    for key in section_order:
        label = _label_for(key)
        heading_text = label if key in NUMBERED_KEYS_EXCLUDED else f"{number}. {label}"
        if key not in NUMBERED_KEYS_EXCLUDED:
            number += 1
        _add_heading(document, heading_text, body_font)
        _add_paragraphs(document, sections.get(key, []), missing_text)

    _add_heading(document, ("References" if not is_ko else "References (참고문헌)"), body_font)
    reference_list = draft.get("reference_list") or []
    if not reference_list:
        p = document.add_paragraph()
        p.paragraph_format.line_spacing = 2.0
        run = p.add_run("[NEEDS DATA: 참고문헌이 제공되지 않았습니다]" if is_ko else "[NEEDS DATA: no references provided]")
        run.italic = True
    else:
        for ref in reference_list:
            p = document.add_paragraph()
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.first_line_indent = Inches(-0.3)
            p.add_run(ref)

    document.save(output_path)


def build_checklist_report_docx(
    draft_title: str,
    fact_check: dict,
    language_check: dict,
    static_checks: dict,
    final_verdict: str,
    output_path: str,
    language: str = "en",
) -> None:
    """저자 본인이 참고할 작업용 체크리스트 문서 (저널에 제출하는 문서가 아님)"""
    is_ko = language == "ko"
    report_font = _report_font(language)

    document = Document()
    style = document.styles["Normal"]
    _apply_font(style.font, report_font)
    style.font.size = Pt(11)

    title_p = document.add_paragraph()
    run = title_p.add_run(f"{'체크리스트 리포트' if is_ko else 'Checklist Report'} — {draft_title}")
    _apply_font(run.font, report_font)
    run.bold = True
    run.font.size = Pt(16)

    note_p = document.add_paragraph()
    note_text = (
        "이 문서는 저자 본인을 위한 작업용 문서이며, 저널에 제출하는 파일이 아닙니다. "
        "이 프로그램은 외부 문헌 데이터베이스 대조 표절 검사(Turnitin, iThenticate 등)나 자기표절 검사를 "
        "수행하지 않습니다 — 투고 전 소속 기관의 표절 검사 도구를 반드시 사용하세요."
        if is_ko else
        "This is a working document for the author, not a file to submit to the journal. "
        "This program does not perform plagiarism checks against external literature databases "
        "(e.g., Turnitin, iThenticate) or self-plagiarism checks - use your institution's dedicated "
        "tool for that before submission."
    )
    note_run = note_p.add_run(note_text)
    note_run.italic = True

    def add_section(heading: str):
        p = document.add_paragraph()
        run = p.add_run(heading)
        run.bold = True
        run.font.size = Pt(13)

    def add_bullets(items, empty_text=None):
        empty_text = empty_text or ("(없음)" if is_ko else "(none found)")
        if not items:
            document.add_paragraph(empty_text, style="List Bullet")
            return
        for item in items:
            document.add_paragraph(str(item), style="List Bullet")

    add_section("최종 판정" if is_ko else "Final Verdict")
    document.add_paragraph(final_verdict)

    add_section("기계적 점검" if is_ko else "Mechanical Checks")
    document.add_paragraph(
        f"{'분량' if is_ko else 'Word count'}: {static_checks.get('word_count')} "
        f"({'목표' if is_ko else 'target'} {static_checks.get('word_count_target')}) - {static_checks.get('word_count_verdict')}"
    )
    document.add_paragraph(
        f"{'잔여 [NEEDS DATA] 표시' if is_ko else 'Remaining [NEEDS DATA] placeholders'}: "
        f"{static_checks.get('remaining_needs_data_count')}"
    )
    add_bullets(static_checks.get("remaining_needs_data_samples", []))
    cc = static_checks.get("citation_check", {})
    document.add_paragraph(
        f"{'참고문헌' if is_ko else 'References'}: {cc.get('reference_count', 0)}, "
        f"{'본문 인용 수' if is_ko else 'in-text citations found'}: {cc.get('in_text_citation_count', 0)}"
    )
    document.add_paragraph("본문에 인용되지 않은 참고문헌:" if is_ko else "References never cited in text:")
    add_bullets(cc.get("references_never_cited_in_text", []))
    document.add_paragraph("참고문헌 목록에 없는 본문 인용:" if is_ko else "In-text citations missing from reference list:")
    add_bullets(cc.get("in_text_citations_missing_from_reference_list", []))
    document.add_paragraph(
        f"{'윤리 정보 제공 여부' if is_ko else 'Ethics statement provided'}: {static_checks.get('has_ethics_statement')}"
    )

    add_section("팩트체크 (Phase 3)" if is_ko else "Fact Check (Phase 3)")
    document.add_paragraph(fact_check.get("overall_note", ""))
    document.add_paragraph("근거 부족 주장:" if is_ko else "Unsupported claims:")
    add_bullets([f"[{c.get('section')}] {c.get('excerpt')} — {c.get('issue')}" for c in fact_check.get("unsupported_claims", [])])
    document.add_paragraph("재현성 이슈:" if is_ko else "Reproducibility issues:")
    add_bullets([f"{c.get('excerpt')} — {c.get('issue')}" for c in fact_check.get("reproducibility_issues", [])])

    add_section("언어 검토 (Phase 3)" if is_ko else "Language Check (Phase 3)")
    document.add_paragraph(language_check.get("overall_note", ""))
    document.add_paragraph("어색한 문장:" if is_ko else "Awkward sentences:")
    add_bullets([f"[{c.get('section')}] {c.get('issue')}" for c in language_check.get("awkward_sentences", [])])
    document.add_paragraph("약어 이슈:" if is_ko else "Abbreviation issues:")
    add_bullets([f"{c.get('abbreviation')}: {c.get('issue')}" for c in language_check.get("abbreviation_issues", [])])

    document.save(output_path)


def build_reviewer_qa_docx(draft_title: str, review_qa: dict, output_path: str, language: str = "en") -> None:
    """예상 심사위원 질문 + 디펜스 문서 (저자용 대비 자료, 제출용 아님)"""
    is_ko = language == "ko"
    report_font = _report_font(language)

    document = Document()
    style = document.styles["Normal"]
    _apply_font(style.font, report_font)
    style.font.size = Pt(11)

    title_p = document.add_paragraph()
    run = title_p.add_run(f"{'예상 심사위원 Q&A' if is_ko else 'Anticipated Reviewer Q&A'} — {draft_title}")
    _apply_font(run.font, report_font)
    run.bold = True
    run.font.size = Pt(16)

    note_p = document.add_paragraph()
    note_text = (
        "이 문서는 저자용 대비 자료(모의 디펜스, 공동저자 논의용)이며, 저널에 제출하는 파일이 아닙니다. "
        "실제로 방어할 수 없는 한계가 드러난 질문의 경우, 디펜스는 억지 반박을 지어내지 않고 정직하게 "
        "한계를 인정하도록 작성되어 있습니다 — 실제 리뷰어 답변으로 옮길 때는 반드시 직접 검토하세요."
        if is_ko else
        "This is a preparation document for the authors (e.g., for a mock defense or co-author "
        "discussion), not a file to submit to the journal. Where a question exposes a genuine "
        "unaddressed limitation, the suggested defense honestly concedes it rather than fabricating "
        "a rebuttal - do not alter these into false rebuttals when actually responding to reviewers."
    )
    note_run = note_p.add_run(note_text)
    note_run.italic = True

    questions = review_qa.get("questions", [])
    if not questions:
        document.add_paragraph("생성된 질문이 없습니다." if is_ko else "No questions were generated.")
        document.save(output_path)
        return

    major = [q for q in questions if q.get("severity") == "major"]
    minor = [q for q in questions if q.get("severity") != "major"]
    summary_p = document.add_paragraph()
    summary_text = (
        f"예상 질문 {len(questions)}건 (주요 {len(major)}건, 경미 {len(minor)}건)"
        if is_ko else
        f"{len(questions)} anticipated questions ({len(major)} major, {len(minor)} minor)"
    )
    summary_run = summary_p.add_run(summary_text)
    summary_run.bold = True

    for i, q in enumerate(questions, 1):
        heading_p = document.add_paragraph()
        heading_p.paragraph_format.space_before = Pt(14)
        sev = (q.get("severity") or "minor").upper()
        run = heading_p.add_run(f"Q{i}. [{q.get('category', 'other').upper()} · {sev}] {q.get('question', '')}")
        run.bold = True
        run.font.size = Pt(12)

        defense_p = document.add_paragraph()
        defense_p.paragraph_format.left_indent = Inches(0.25)
        label_run = defense_p.add_run("디펜스: " if is_ko else "Suggested defense: ")
        label_run.italic = True
        label_run.bold = True
        defense_p.add_run(q.get("defense", ""))

    document.save(output_path)
