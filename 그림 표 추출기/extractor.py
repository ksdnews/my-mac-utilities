import os
import sys
import zipfile
import xml.etree.ElementTree as ET
import traceback
import tempfile
import shutil
import unicodedata

# XML Namespaces for HWPX parsing & generation
NAMESPACES = {
    'hh': 'http://www.hancom.co.kr/hwpml/2011/head',
    'hp': 'http://www.hancom.co.kr/hwpml/2011/paragraph',
    'hs': 'http://www.hancom.co.kr/hwpml/2011/core',
    'hc': 'http://www.hancom.co.kr/hwpml/2011/container',
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
}

def extract_docx(file_path, temp_dir):
    """Extracts images and tables from a .docx file"""
    images = []
    tables = []
    
    # 1. Extract Images directly from Zip structure (most reliable)
    with zipfile.ZipFile(file_path, 'r') as z:
        for f in z.namelist():
            if f.startswith('word/media/'):
                base_name = os.path.basename(f)
                out_path = os.path.join(temp_dir, base_name)
                with open(out_path, 'wb') as img_out:
                    img_out.write(z.read(f))
                images.append(out_path)
                
    # 2. Extract Tables using python-docx
    try:
        import docx
        doc = docx.Document(file_path)
        for idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                tables.append(table_data)
    except Exception as e:
        print(f"[Warning] Failed to parse Word tables via python-docx: {e}")
        
    return images, tables

def extract_pdf_fallback(file_path, temp_dir):
    """Fallback method to extract images and tables from a .pdf file using pdfplumber and pypdf"""
    images = []
    tables = []
    
    # 1. Extract Tables via pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                extracted = page.extract_tables()
                for table in extracted:
                    cleaned_table = []
                    for row in table:
                        if not row:
                            continue
                        cleaned_row = [str(cell).strip() if cell is not None else "" for cell in row]
                        cleaned_table.append(cleaned_row)
                    if cleaned_table:
                        tables.append(cleaned_table)
    except Exception as e:
        print(f"[Warning] Failed to extract tables via pdfplumber: {e}")
        
    # 2. Extract Images via pypdf (fallback to pdfplumber if pypdf unavailable)
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        img_counter = 1
        for page_idx, page in enumerate(reader.pages):
            for image_file_object in page.images:
                img_name = f"pdf_p{page_idx+1}_{img_counter}_{image_file_object.name}"
                out_path = os.path.join(temp_dir, img_name)
                with open(out_path, "wb") as fp:
                    fp.write(image_file_object.data)
                images.append(out_path)
                img_counter += 1
    except Exception as e:
        print(f"[Warning] Failed to extract PDF images via pypdf: {e}")
        # Try pdfplumber image extraction as fallback
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page_idx, page in enumerate(pdf.pages):
                    for img_idx, img in enumerate(page.images):
                        # Decode image bounding box
                        bbox = (img["x0"], img["top"], img["x1"], img["bottom"])
                        cropped = page.within_bbox(bbox)
                        # Export high-res image
                        img_obj = cropped.to_image(resolution=150)
                        img_name = f"pdf_p{page_idx+1}_img{img_idx+1}.png"
                        out_path = os.path.join(temp_dir, img_name)
                        img_obj.save(out_path, format="PNG")
                        images.append(out_path)
        except Exception as e2:
            print(f"[Warning] Fallback PDF image extraction failed: {e2}")

    return images, tables

def extract_pdf(file_path, temp_dir):
    """Extracts images and tables from a .pdf file using PyMuPDF (fitz) with smart filtering and page rendering"""
    try:
        import fitz
    except ImportError:
        print("[Warning] fitz (PyMuPDF) is not installed. Falling back to pdfplumber/pypdf.")
        return extract_pdf_fallback(file_path, temp_dir)
        
    images = []
    tables = []
    
    try:
        doc = fitz.open(file_path)
        
        for page_idx, page in enumerate(doc):
            # --- 1. 표(Table) 추출 ---
            try:
                tables_finder = page.find_tables()
                if tables_finder and tables_finder.tables:
                    for t in tables_finder.tables:
                        # 격자 좌표 리스트 구성
                        x_set = set()
                        y_set = set()
                        for cell in t.cells:
                            if cell is not None:
                                x_set.add(cell[0])
                                x_set.add(cell[2])
                                y_set.add(cell[1])
                                y_set.add(cell[3])
                        
                        x_coords = sorted(list(x_set))
                        y_coords = sorted(list(y_set))
                        
                        raw_matrix = t.extract()
                        
                        # 격자 인덱스를 구하는 보조 함수
                        def get_coord_index(val, coord_list):
                            closest_idx = 0
                            min_diff = float('inf')
                            for idx, c_val in enumerate(coord_list):
                                diff = abs(c_val - val)
                                if diff < min_diff:
                                    min_diff = diff
                                    closest_idx = idx
                            return closest_idx
                            
                        cells_info = []
                        for cell in t.cells:
                            if cell is None:
                                continue
                            x0, y0, x1, y1 = cell
                            c_start = get_coord_index(x0, x_coords)
                            c_end = get_coord_index(x1, x_coords)
                            r_start = get_coord_index(y0, y_coords)
                            r_end = get_coord_index(y1, y_coords)
                            
                            cell_text = ""
                            if r_start < len(raw_matrix) and c_start < len(raw_matrix[r_start]):
                                cell_text = raw_matrix[r_start][c_start]
                                if cell_text is None:
                                    cell_text = ""
                                    
                            cells_info.append({
                                'r_start': r_start,
                                'r_end': r_end,
                                'c_start': c_start,
                                'c_end': c_end,
                                'text': cell_text
                            })
                            
                        # 병합 정보를 포함한 딕셔너리로 저장
                        tables.append({
                            'type': 'merged_table',
                            'row_count': t.row_count,
                            'col_count': t.col_count,
                            'cells': cells_info
                        })
            except Exception as te:
                print(f"[Warning] Failed to extract tables from page {page_idx+1} using PyMuPDF: {te}")
            
            # --- 2. 그림(Image) 추출 ---
            text = page.get_text().strip()
            image_list = page.get_images(full=True)
            
            # 텍스트가 적고(150자 미만) 이미지가 있는 전면 그래픽/포스터 페이지는 표 모양의 상자를 제외하고 알맹이 그림만 추출
            if len(text) < 150 and len(image_list) > 0:
                try:
                    union_rect = None
                    for img_info in image_list:
                        xref = img_info[0]
                        rects = page.get_image_rects(xref)
                        for r in rects:
                            if r.x1 - r.x0 >= 20 and r.y1 - r.y0 >= 20:
                                if union_rect is None:
                                    union_rect = fitz.Rect(r)
                                else:
                                    union_rect = union_rect | r
                    
                    # 큰 수직 테두리선(드로잉 경로)을 분석하여 포스터의 상자 영역을 찾음
                    left_x, right_x = None, None
                    top_y, bottom_y = None, None
                    
                    drawings = page.get_drawings()
                    for d in drawings:
                        rect = d["rect"]
                        width = rect.x1 - rect.x0
                        height = rect.y1 - rect.y0
                        # 수직선 검출 (가로폭이 아주 얇고, 높이가 300포인트 이상인 테두리선)
                        if width <= 1.0 and height >= 300:
                            # 좌측 테두리선 후보 (x 좌표가 50~70 사이)
                            if 50 <= rect.x0 <= 70:
                                left_x = rect.x0
                                top_y = rect.y0 if top_y is None else min(top_y, rect.y0)
                                bottom_y = rect.y1 if bottom_y is None else max(bottom_y, rect.y1)
                            # 우측 테두리선 후보 (x 좌표가 520~550 사이)
                            elif 520 <= rect.x0 <= 550:
                                right_x = rect.x0
                                top_y = rect.y0 if top_y is None else min(top_y, rect.y0)
                                bottom_y = rect.y1 if bottom_y is None else max(bottom_y, rect.y1)
                                
                    # 만약 좌우 수직 테두리선과 상하 범위를 모두 찾아냈다면 그 영역을 포스터 상자로 정의
                    if left_x is not None and right_x is not None and top_y is not None and bottom_y is not None:
                        # 외곽 검은 테두리선(1pt 두께)을 안쪽으로 1.5포인트 깎아 선까지 완벽히 지우고 안쪽 콘텐츠만 크롭
                        clip_rect = fitz.Rect(left_x + 1.5, top_y + 1.5, right_x - 1.5, bottom_y - 1.5)
                        pix = page.get_pixmap(clip=clip_rect, dpi=150)
                    else:
                        # 찾지 못한 경우 이미지들의 union_rect 영역 사용
                        if union_rect and not union_rect.is_empty:
                            padding = 5
                            x0 = max(0, union_rect.x0 - padding)
                            y0 = max(0, union_rect.y0 - padding)
                            x1 = min(page.rect.x1, union_rect.x1 + padding)
                            y1 = min(page.rect.y1, union_rect.y1 + padding)
                            clip_rect = fitz.Rect(x0, y0, x1, y1)
                            pix = page.get_pixmap(clip=clip_rect, dpi=150)
                        else:
                            pix = page.get_pixmap(dpi=150)
                        
                    img_name = f"pdf_page_{page_idx+1}_render.png"
                    out_path = os.path.join(temp_dir, img_name)
                    pix.save(out_path)
                    images.append(out_path)
                    continue  # 전체 페이지 렌더링을 수행했으므로 조각 이미지 추출은 스킵
                except Exception as re:
                    print(f"[Warning] Failed to render page {page_idx+1} area: {re}")
            
            # 일반 텍스트 페이지는 유의미한 개별 이미지 객체만 추출 (잘림 방지를 위해 바운딩 박스 크롭 렌더링 적용)
            for img_idx, img_info in enumerate(image_list):
                try:
                    xref = img_info[0]
                    rects = page.get_image_rects(xref)
                    
                    if rects:
                        # 페이지 상의 실제 이미지 사각형 영역들을 클리핑하여 렌더링 (마스크 결합 및 잘림 방지)
                        for r_idx, rect in enumerate(rects):
                            width_pt = rect.x1 - rect.x0
                            height_pt = rect.y1 - rect.y0
                            
                            # 가로/세로가 약 20pt(약 40px) 이상인 유의미한 크기만 추출 (로고 등 얇은 배너 지원)
                            if width_pt >= 20 and height_pt >= 20:
                                aspect_ratio = width_pt / height_pt if height_pt > 0 else 0
                                if 0.1 <= aspect_ratio <= 10.0:
                                    # 고해상도(150 DPI)로 해당 영역을 클리핑 렌더링하여 PNG로 저장
                                    pix = page.get_pixmap(clip=rect, dpi=150)
                                    suffix = f"_{r_idx}" if len(rects) > 1 else ""
                                    img_name = f"pdf_p{page_idx+1}_img{img_idx+1}_{xref}{suffix}.png"
                                    out_path = os.path.join(temp_dir, img_name)
                                    pix.save(out_path)
                                    images.append(out_path)
                    else:
                        # 렌더링 영역 정보를 얻지 못할 경우에만 원본 데이터를 덤프하고 Pillow로 표준화
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        width = base_image.get("width", 0)
                        height = base_image.get("height", 0)
                        size = len(image_bytes)
                        
                        if width >= 40 and height >= 40 and size >= 2048:
                            aspect_ratio = width / height if height > 0 else 0
                            if 0.1 <= aspect_ratio <= 10.0:
                                import io
                                from PIL import Image as PILImage
                                # Pillow를 사용해 열어서 표준 PNG 포맷으로 파일 저장 (삽입 실패 에러 완벽 방지)
                                img = PILImage.open(io.BytesIO(image_bytes))
                                img_name = f"pdf_p{page_idx+1}_img{img_idx+1}_{xref}.png"
                                out_path = os.path.join(temp_dir, img_name)
                                img.save(out_path, format="PNG")
                                images.append(out_path)
                except Exception as ie:
                    print(f"[Warning] Failed to extract image {img_idx+1} on page {page_idx+1}: {ie}")
                    
    except Exception as e:
        print(f"[Error] Failed to process PDF via PyMuPDF: {e}")
        # 예외 발생 시에도 fallback으로 우회 시도
        return extract_pdf_fallback(file_path, temp_dir)
        
    return images, tables

def extract_hwpx(file_path, temp_dir):
    """Extracts images and tables from a .hwpx file"""
    images = []
    tables = []
    
    # 1. Extract Images from BinData
    with zipfile.ZipFile(file_path, 'r') as z:
        for f in z.namelist():
            if f.startswith('BinData/'):
                base_name = os.path.basename(f)
                if base_name:
                    out_path = os.path.join(temp_dir, base_name)
                    with open(out_path, 'wb') as img_out:
                        img_out.write(z.read(f))
                    images.append(out_path)
                    
        # 2. Extract Tables from Contents/section0.xml
        try:
            sec_file = None
            for f in z.namelist():
                if f.endswith('section0.xml'):
                    sec_file = f
                    break
            if sec_file:
                xml_content = z.read(sec_file)
                root = ET.fromstring(xml_content)
                # Register namespaces locally
                for prefix, uri in NAMESPACES.items():
                    ET.register_namespace(prefix, uri)
                
                # Find all tbl elements: hp:tbl
                tbl_query = './/{http://www.hancom.co.kr/hwpml/2011/paragraph}tbl'
                for tbl in root.findall(tbl_query):
                    table_data = []
                    # Find rows: hp:tr
                    for tr in tbl.findall('.//{http://www.hancom.co.kr/hwpml/2011/paragraph}tr'):
                        row_data = []
                        # Find cells: hp:tc
                        for tc in tr.findall('.//{http://www.hancom.co.kr/hwpml/2011/paragraph}tc'):
                            # Accumulate text inside cell
                            cell_text = ""
                            for t in tc.findall('.//{http://www.hancom.co.kr/hwpml/2011/paragraph}t'):
                                if t.text:
                                    cell_text += t.text
                            row_data.append(cell_text.strip())
                        table_data.append(row_data)
                    if table_data:
                        tables.append(table_data)
        except Exception as e:
            print(f"[Warning] Failed to parse HWPX XML tables: {e}")
            
    return images, tables

def extract_hwp(file_path, temp_dir):
    """Extracts images and tables from a legacy binary .hwp file using olefile"""
    images = []
    tables = []
    
    # 1. Extract internal images via OLE storage
    try:
        import olefile
        if olefile.isOleFile(file_path):
            ole = olefile.OleFileIO(file_path)
            # Find and extract streams under BinData
            for stream in ole.listdir():
                if len(stream) > 1 and stream[0] == 'BinData':
                    stream_path = '/'.join(stream)
                    # Deduce image extension
                    ext = "png"
                    lower_name = stream[1].lower()
                    if "jpg" in lower_name or "jpeg" in lower_name:
                        ext = "jpg"
                    elif "gif" in lower_name:
                        ext = "gif"
                    elif "bmp" in lower_name:
                        ext = "bmp"
                        
                    # Output image file name
                    base_name = f"hwp_{stream[1]}.{ext}"
                    out_path = os.path.join(temp_dir, base_name)
                    with open(out_path, 'wb') as img_out:
                        img_out.write(ole.openstream(stream_path).read())
                    images.append(out_path)
                    
            # 2. Parse Text / Table approximation via zlib decompressed BodyText stream
            try:
                text_content = ""
                # HWP stores body text in streams named 'BodyText/Section0', 'BodyText/Section1', etc.
                for stream in ole.listdir():
                    if len(stream) > 1 and stream[0] == 'BodyText':
                        stream_path = '/'.join(stream)
                        compressed_data = ole.openstream(stream_path).read()
                        
                        # Decompress HWP stream (using raw zlib decompression)
                        import zlib
                        try:
                            # HWP compression skips zlib headers sometimes, try wbits=-15 (raw deflate) or automatic detection
                            decompressed = zlib.decompress(compressed_data, -15)
                        except Exception:
                            try:
                                decompressed = zlib.decompress(compressed_data)
                            except Exception:
                                decompressed = compressed_data
                                
                        # Scan printable text characters from decompressed binary stream (rough parser)
                        # HWP uses 2-byte characters (Unicode). Filter out control characters.
                        chars = []
                        i = 0
                        while i < len(decompressed) - 1:
                            char_code = decompressed[i] + (decompressed[i+1] << 8)
                            # Unicode range check for Korean and English
                            if (0xac00 <= char_code <= 0xd7a3) or (32 <= char_code <= 126) or char_code == 10 or char_code == 13:
                                chars.append(chr(char_code))
                            i += 2
                        text_content += "".join(chars)
                
                # Check if there are lines resembling table structures (e.g. columns separated by multiple spaces or tabs)
                # Since precise table parsing in binary HWP is highly difficult without a complete renderer,
                # we group paragraphs that resemble lists/tables and store them as Text Tables.
                lines = [line.strip() for line in text_content.split('\n') if line.strip()]
                current_table = []
                for line in lines:
                    # Rough heuristic: if line contains '|' or '\t' or multiple spaces, treat as row
                    parts = []
                    if '\t' in line:
                        parts = [p.strip() for p in line.split('\t') if p.strip()]
                    elif '   ' in line:
                        parts = [p.strip() for p in line.split('   ') if p.strip()]
                    
                    if len(parts) >= 2:
                        current_table.append(parts)
                    else:
                        if len(current_table) >= 2:
                            tables.append(current_table)
                            current_table = []
                if len(current_table) >= 2:
                    tables.append(current_table)
            except Exception as pe:
                print(f"[Warning] Failed to scan legacy HWP text streams: {pe}")
            
            ole.close()
    except Exception as e:
        print(f"[Warning] Legacy HWP OLE parsing failed: {e}")
        
    return images, tables

def build_docx_output(images, tables, output_path, input_file):
    """Dynamically packages extracted tables and images into a valid DOCX file and saves raw images to a directory"""
    import docx
    from docx.shared import Inches
    
    # Normalize paths to NFC (composes Korean jamo on macOS)
    output_path = unicodedata.normalize('NFC', output_path)
    input_file = unicodedata.normalize('NFC', input_file)
    
    # Get original filename without extension (기존 파일 이름)
    input_base = os.path.splitext(os.path.basename(input_file))[0]
    
    # 1. Create directory for raw images next to output file
    base_no_ext = os.path.splitext(output_path)[0]
    img_save_dir = f"{base_no_ext}_추출그림"
    
    if images:
        os.makedirs(img_save_dir, exist_ok=True)
        for idx, img_path in enumerate(images):
            ext = os.path.splitext(img_path)[1].lower()
            # New filename format: "기존 파일 이름_그림 추출 [번호].[확장자]"
            new_name = f"{input_base}_그림 추출 {idx+1}{ext}"
            dest_path = os.path.join(img_save_dir, new_name)
            shutil.copy2(img_path, dest_path)
        print(f"📁 개별 그림 파일들이 다음 폴더에 저장되었습니다: {img_save_dir}")
    
    # 2. Build the DOCX document
    doc = docx.Document()
    
    # Add title and headers
    doc.add_heading(unicodedata.normalize('NFC', "문서 이미지 & 표 자동 추출 결과 리포트"), level=1)
    doc.add_paragraph(unicodedata.normalize('NFC', "본 문서는 원래 문서로부터 그림 파일과 표 데이터를 자동으로 분리해 저장한 결과물입니다."))
    doc.add_paragraph()
    
    # Add Tables
    if tables:
        doc.add_heading(unicodedata.normalize('NFC', f"추출된 표 목록 - 총 {len(tables)}개"), level=2)
        for t_idx, table in enumerate(tables):
            doc.add_heading(unicodedata.normalize('NFC', f"표 #{t_idx + 1}"), level=3)
            if not table:
                doc.add_paragraph(unicodedata.normalize('NFC', "(빈 표 데이터)"))
                continue
            
            # 딕셔너리 형태의 병합 표인 경우
            if isinstance(table, dict) and table.get('type') == 'merged_table':
                num_rows = table['row_count']
                num_cols = table['col_count']
                
                if num_rows > 0 and num_cols > 0:
                    docx_table = doc.add_table(rows=num_rows, cols=num_cols)
                    docx_table.style = 'Table Grid'
                    
                    # 1단계: 셀 병합 적용
                    for cell_info in table['cells']:
                        r_start = cell_info['r_start']
                        r_end = cell_info['r_end']
                        c_start = cell_info['c_start']
                        c_end = cell_info['c_end']
                        
                        if (r_end - r_start > 1) or (c_end - c_start > 1):
                            try:
                                cell_a = docx_table.cell(r_start, c_start)
                                cell_b = docx_table.cell(r_end - 1, c_end - 1)
                                cell_a.merge(cell_b)
                            except Exception as me:
                                print(f"[Warning] Table cell merge failed: {me}")
                                
                    # 2단계: 텍스트 채우기 (정규화 포함)
                    for cell_info in table['cells']:
                        r_start = cell_info['r_start']
                        c_start = cell_info['c_start']
                        cell_text = cell_info['text']
                        if cell_text:
                            try:
                                docx_table.cell(r_start, c_start).text = unicodedata.normalize('NFC', str(cell_text))
                            except Exception as te:
                                print(f"[Warning] Failed to set text on table cell ({r_start}, {c_start}): {te}")
            else:
                # 기존 2차원 리스트 표 처리 (하위 호환성)
                num_rows = len(table)
                num_cols = max(len(row) for row in table) if table else 0
                
                if num_rows > 0 and num_cols > 0:
                    docx_table = doc.add_table(rows=num_rows, cols=num_cols)
                    docx_table.style = 'Table Grid'
                    
                    for r_idx, row in enumerate(table):
                        for c_idx, val in enumerate(row):
                            if c_idx < len(docx_table.rows[r_idx].cells):
                                cell_val = str(val) if val is not None else ""
                                docx_table.rows[r_idx].cells[c_idx].text = unicodedata.normalize('NFC', cell_val)
            doc.add_paragraph()
    else:
        doc.add_heading(unicodedata.normalize('NFC', "추출된 표 목록"), level=2)
        doc.add_paragraph(unicodedata.normalize('NFC', "인식된 표 데이터가 없습니다."))
        doc.add_paragraph()
        
    # Add Images to document
    if images:
        doc.add_heading(unicodedata.normalize('NFC', f"추출된 그림 목록 - 총 {len(images)}개"), level=2)
        for i_idx, img_path in enumerate(images):
            ext = os.path.splitext(img_path)[1].lower()
            new_name = f"{input_base}_그림 추출 {i_idx+1}{ext}"
            doc.add_paragraph(unicodedata.normalize('NFC', f"📷 그림 #{i_idx + 1} (파일명: {new_name})"))
            
            try:
                # Add picture with fixed width, keeping aspect ratio
                doc.add_picture(img_path, width=Inches(5))
            except Exception as ie:
                doc.add_paragraph(unicodedata.normalize('NFC', f"[그림 삽입 실패: {ie}]"))
            doc.add_paragraph()
    else:
        doc.add_heading(unicodedata.normalize('NFC', "추출된 그림 목록"), level=2)
        doc.add_paragraph(unicodedata.normalize('NFC', "추출된 이미지 파일이 없습니다."))
        doc.add_paragraph()
        
    # Save the document
    doc.save(output_path)

def main():
    if len(sys.argv) < 3:
        print("Usage: python extractor.py <input_file> <output_docx> [extract_images: True/False] [extract_tables: True/False]")
        sys.exit(1)
        
    input_file = sys.argv[1]
    output_docx = sys.argv[2]
    
    extract_images = True
    extract_tables = True
    
    if len(sys.argv) > 3:
        extract_images = sys.argv[3].lower() == 'true'
    if len(sys.argv) > 4:
        extract_tables = sys.argv[4].lower() == 'true'
    
    if not os.path.exists(input_file):
        print(f"Error: Input file '{input_file}' not found.")
        sys.exit(1)
        
    ext = os.path.splitext(input_file)[1].lower()
    
    # Prepare temporary directory for image caching
    temp_dir = tempfile.mkdtemp()
    
    try:
        print(f"[{ext.upper()}] 파싱을 진행하고 있습니다...")
        images = []
        tables = []
        
        if ext == '.docx' or ext == '.doc':
            images, tables = extract_docx(input_file, temp_dir)
        elif ext == '.pdf':
            images, tables = extract_pdf(input_file, temp_dir)
        elif ext == '.hwpx':
            images, tables = extract_hwpx(input_file, temp_dir)
        elif ext == '.hwp':
            images, tables = extract_hwp(input_file, temp_dir)
        else:
            print(f"Error: Unsupported file format '{ext}'.")
            sys.exit(1)
            
        # Filter based on user preferences
        if not extract_images:
            images = []
        if not extract_tables:
            tables = []
            
        print(f"🔍 추출 통계: 그림 {len(images)}개, 표 {len(tables)}개 탑지됨.")
        
        # Build packaged DOCX document
        print("💾 추출 데이터를 바탕으로 DOCX 결과 리포트를 작성하는 중...")
        build_docx_output(images, tables, output_docx, input_file)
        print(f"🎉 성공: DOCX 파일 작성 완료 -> {output_docx}")
        
    except Exception as e:
        print(f"Error during extraction: {e}")
        traceback.print_exc()
        sys.exit(1)
    finally:
        # Clean up temporary folders
        shutil.rmtree(temp_dir)

if __name__ == '__main__':
    main()
